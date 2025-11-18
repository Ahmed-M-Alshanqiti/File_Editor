import os
import subprocess
from decimal import Decimal
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import FileResponse, Http404, HttpResponseForbidden, HttpResponse
from django.conf import settings
from django.contrib import messages
from django.views.decorators.clickjacking import xframe_options_exempt
from django.core.files import File
from django.db import transaction
from django.utils import timezone
from django.contrib.auth.models import User

from users.models import Profile
from .models import (
    UploadedFile,
    Comment,
    UploadedFileVersion,
    Notification,
    FileStatus,
    ChangeTypes,
)
from .forms import UploadFileForm, CommentForm

# Optional: docx / excel support
try:
    from docx import Document
    DOCX_SUPPORTED = True
except Exception:
    DOCX_SUPPORTED = False

try:
    from openpyxl import load_workbook
    EXCEL_SUPPORTED = True
except Exception:
    EXCEL_SUPPORTED = False


# -------------------------
# Helper: extension checks
# -------------------------
TEXT_EXTENSIONS = {'.txt', '.md', '.py', '.json', '.csv', '.html', '.css', '.js'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp'}
PDF_EXTENSIONS = {'.pdf'}
DOCX_EXTENSIONS = {'.docx'}
EXCEL_EXTENSIONS = {'.xlsx'}


# -------------------------
# Role / notification helpers
# -------------------------
def _get_profile(user):
    if not user.is_authenticated:
        return None
    try:
        return user.profile
    except Profile.DoesNotExist:
        return None


def is_program_super_user(user):
    profile = _get_profile(user)
    return bool(profile and profile.role == Profile.Roles.SUPER_REVIEWER)


def is_auditor(user):
    profile = _get_profile(user)
    return bool(profile and profile.role == Profile.Roles.AUDITOR)


def can_upload_files(user):
    return user.is_authenticated and (is_program_super_user(user) or is_auditor(user))


def notify_users(user_qs, sender, notif_type, message, file_obj):
    for recipient in user_qs:
        Notification.objects.create(
            recipient=recipient,
            sender=sender,
            notification_type=notif_type,
            message=message,
            related_file=file_obj,
        )


def notify_super_reviewers(file_obj, sender, notif_type, message):
    reviewers = User.objects.filter(profile__role=Profile.Roles.SUPER_REVIEWER, is_active=True)
    notify_users(reviewers, sender, notif_type, message, file_obj)


# -------------------------
# List / Upload / Detail
# -------------------------
@login_required
def file_list(request):
    files = UploadedFile.objects.all().order_by('-uploaded_at')
    return render(request, 'file_list.html', {'files': files, 'FileStatus': FileStatus})


@login_required
def file_upload(request):
    if not can_upload_files(request.user):
        return HttpResponseForbidden("You are not allowed to upload files.")

    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            file_inst = form.save(commit=False)
            file_inst.owner = request.user
            file_inst.status = FileStatus.PENDING
            file_inst.version_number = Decimal('1.0')
            file_inst.reviewed_at = None
            file_inst.reviewed_by = None
            file_inst.save()
            UploadedFileVersion.objects.create(
                file=file_inst,
                version_label=file_inst.version_label,
                change_type=ChangeTypes.MAJOR,
                comment="Initial upload",
                created_by=request.user,
            )

            notify_super_reviewers(
                file_inst,
                request.user,
                Notification.Types.FILE_SUBMITTED,
                f"{request.user.username} uploaded {file_inst.filename} (version {file_inst.version_label}).",
                file_inst,
            )

            messages.success(request, "File uploaded.")
            return redirect('file_list')
    else:
        form = UploadFileForm()
    return render(request, 'file_upload.html', {'form': form})


def file_detail(request, pk):
    """
    Standard file detail view - shows file info, comments, and PDF preview
    """
    file_obj = get_object_or_404(UploadedFile, pk=pk)
    
    comments = file_obj.comments.all().order_by('-created_at')
    comment_form = CommentForm() if request.user.is_authenticated and request.user == file_obj.owner else None
    versions = file_obj.versions.select_related('created_by')

    can_download_original = request.user.is_authenticated and request.user == file_obj.owner
    can_review = is_program_super_user(request.user)

    return render(request, 'file_detail.html', {
        'file': file_obj,
        'comments': comments,
        'comment_form': comment_form,
        'versions': versions,
        'can_download_original': can_download_original,
        'can_review': can_review,
        'FileStatus': FileStatus,
    })


# -------------------------
# PDF Serving Views
# -------------------------
from django.views.decorators.csrf import csrf_exempt

@login_required
@xframe_options_exempt
def view_pdf(request, pk):
    """
    Very simple PDF-serving view used by PDF.js.

    Always returns either:
      - 200 with the PDF stream, or
      - 404 if the converted PDF does not exist.
    """
    from django.http import FileResponse

    file_obj = get_object_or_404(UploadedFile, pk=pk)

    # Expect a converted PDF to be present
    if not file_obj.converted:
        raise Http404("PDF not yet converted")

    pdf_path = file_obj.converted.path

    if not os.path.exists(pdf_path):
        raise Http404("PDF file not found on server")

    # Let Django stream the file; PDF.js works fine with this.
    response = FileResponse(open(pdf_path, "rb"), content_type="application/pdf")
    response["Content-Disposition"] = (
        f'inline; filename="{file_obj.file_name_if_converted or "document.pdf"}"'
    )
    return response


@login_required
def download_pdf(request, pk):
    """
    Force download the converted PDF
    """
    file_obj = get_object_or_404(UploadedFile, pk=pk)
    
    if not file_obj.converted:
        raise Http404("PDF not found. Convert first.")
    
    try:
        pdf_path = file_obj.converted.path
        if not os.path.exists(pdf_path):
            raise Http404("PDF file not found on server")
        
        response = FileResponse(
            open(pdf_path, 'rb'),
            content_type='application/pdf'
        )
        # Force download
        response['Content-Disposition'] = f'attachment; filename="{file_obj.file_name_if_converted or "download.pdf"}"'
        return response
        
    except Exception as e:
        raise Http404(f"Error: {str(e)}")


@login_required
def download_original(request, pk):
    file_obj = get_object_or_404(UploadedFile, pk=pk)
    if request.user != file_obj.owner:
        return HttpResponseForbidden("Only the uploader can download the original file.")

    file_path = file_obj.file.path
    if not os.path.exists(file_path):
        raise Http404("Original file not found on server.")

    response = FileResponse(open(file_path, 'rb'), content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{file_obj.filename}"'
    return response


# -------------------------
# Edit view — full support
# -------------------------
@login_required
def file_edit(request, file_id):
    """
    - Shows preview depending on file extension.
    - Allows in-place editing for TEXT_EXTENSIONS via textarea (editable_text).
    - Allows owner to replace file via UploadFileForm.
    """
    file_obj = get_object_or_404(UploadedFile, pk=file_id)

    if request.user != file_obj.owner:
        return HttpResponseForbidden("You are not allowed to edit this file.")

    # derive extension and path
    filename = file_obj.file.name
    extension = os.path.splitext(filename)[1].lower()
    file_path = file_obj.file.path

    # preview flags and content
    editable_text = None
    text_preview = None
    image_preview = False
    pdf_preview = False

    # prepare previews/content
    if extension in TEXT_EXTENSIONS:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                editable_text = f.read()
                text_preview = editable_text
        except Exception as e:
            text_preview = f"Unable to read file: {e}"

    elif extension in IMAGE_EXTENSIONS:
        image_preview = True

    elif extension in PDF_EXTENSIONS:
        pdf_preview = True

    elif extension in DOCX_EXTENSIONS and DOCX_SUPPORTED:
        # Basic DOCX editing: flatten paragraphs into text
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            editable_text = "\n\n".join(paragraphs)
            text_preview = editable_text
        except Exception as e:
            text_preview = f"Unable to read DOCX content: {e}"
            editable_text = None

    elif extension in EXCEL_EXTENSIONS and EXCEL_SUPPORTED:
        # Basic Excel editing: render first sheet as CSV-like text
        try:
            wb = load_workbook(file_path)
            ws = wb.active
            rows_as_text = []
            for row in ws.iter_rows():
                values = []
                for cell in row:
                    v = "" if cell.value is None else str(cell.value)
                    # simple escaping for commas
                    if "," in v:
                        v = f"\"{v}\""
                    values.append(v)
                # join cells by comma
                rows_as_text.append(",".join(values))
            editable_text = "\n".join(rows_as_text)
            text_preview = editable_text
        except Exception as e:
            text_preview = f"Unable to read Excel content: {e}"
            editable_text = None

    # POST handling
    if request.method == 'POST':
        change_type = request.POST.get('change_type', ChangeTypes.MINOR)
        if change_type not in dict(ChangeTypes.choices):
            messages.error(request, "Select whether this edit is minor or major.")
            return redirect('file_edit', file_id=file_id)

        reconvert_after_edit = False

        # Handle inline edits for supported types
        if 'edited_text' in request.POST:
            if request.user != file_obj.owner:
                return HttpResponseForbidden("No permission.")

            new_text = request.POST.get('edited_text', '')

            # Optional comment attached to this edit
            edit_comment_text = request.POST.get('edit_comment', '').strip()

            try:
                if extension in TEXT_EXTENSIONS:
                    # Save plain text directly back to the file
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(new_text)

                elif extension in DOCX_EXTENSIONS and DOCX_SUPPORTED:
                    # Rebuild a simple DOCX document from the edited text
                    doc = Document()
                    for block in new_text.split("\n\n"):
                        doc.add_paragraph(block.replace("\r\n", "\n"))
                    doc.save(file_path)

                elif extension in EXCEL_EXTENSIONS and EXCEL_SUPPORTED:
                    # Rebuild an Excel workbook from CSV-like text
                    wb = load_workbook(file_path)
                    ws = wb.active

                    # Clear existing content
                    for row in ws.iter_rows():
                        for cell in row:
                            cell.value = None

                    # Populate from edited text
                    import csv
                    from io import StringIO

                    reader = csv.reader(StringIO(new_text))
                    for r_idx, row in enumerate(reader, start=1):
                        for c_idx, value in enumerate(row, start=1):
                            ws.cell(row=r_idx, column=c_idx, value=value)

                    wb.save(file_path)

                else:
                    messages.error(request, "Inline editing not supported for this file type.")
                    return redirect('file_detail', pk=file_id)

                # Update metadata / versioning
                file_obj.bump_version(change_type)
                file_obj.status = FileStatus.PENDING
                file_obj.reviewed_at = None
                file_obj.reviewed_by = None
                file_obj.save()

                UploadedFileVersion.objects.create(
                    file=file_obj,
                    version_label=file_obj.version_label,
                    change_type=change_type,
                    comment=edit_comment_text,
                    created_by=request.user,
                )

                if edit_comment_text:
                    Comment.objects.create(
                        file=file_obj,
                        user=request.user,
                        text=f"[Version {file_obj.version_label}] {edit_comment_text}",
                    )

                messages.success(request, f"Changes saved (version {file_obj.version_label}).")
                reconvert_after_edit = True

                notify_super_reviewers(
                    file_obj,
                    request.user,
                    Notification.Types.FILE_SUBMITTED,
                    f"{request.user.username} updated {file_obj.filename} to version {file_obj.version_label} ({change_type}).",
                    file_obj,
                )
            except Exception as e:
                messages.error(request, f"Failed to save changes: {e}")

            if reconvert_after_edit and file_obj.converted:
                success, feedback = _convert_with_libreoffice(file_obj)
                if success:
                    messages.success(request, f"PDF regenerated: {feedback}")
                else:
                    messages.warning(request, f"File saved, but PDF conversion failed: {feedback}")

            return redirect('file_detail', pk=file_id)

        # Handle file replacement
        form = UploadFileForm(request.POST, request.FILES, instance=file_obj)
        if form.is_valid():
            # Remove old converted PDF if present
            try:
                if file_obj.converted and file_obj.converted.path and os.path.exists(file_obj.converted.path):
                    os.remove(file_obj.converted.path)
                    file_obj.converted = None
                    file_obj.file_name_if_converted = ''
            except Exception:
                pass

            file_inst = form.save(commit=False)
            file_inst.owner = request.user
            file_inst.status = FileStatus.PENDING
            file_inst.reviewed_at = None
            file_inst.reviewed_by = None
            file_inst.bump_version(change_type)
            file_inst.save()

            note = request.POST.get('edit_comment', '').strip()
            UploadedFileVersion.objects.create(
                file=file_inst,
                version_label=file_inst.version_label,
                change_type=change_type,
                comment=note,
                created_by=request.user,
            )

            if note:
                Comment.objects.create(
                    file=file_inst,
                    user=request.user,
                    text=f"[Version {file_inst.version_label}] {note}",
                )

            notify_super_reviewers(
                file_inst,
                request.user,
                Notification.Types.FILE_SUBMITTED,
                f"{request.user.username} replaced {file_inst.filename} (version {file_inst.version_label}).",
                file_inst,
            )

            messages.success(request, f"File replaced. New version: {file_inst.version_label}.")
            return redirect('file_detail', pk=file_id)
        else:
            messages.error(request, "Upload failed. Check file and try again.")

    else:
        form = UploadFileForm(instance=file_obj)

    return render(request, 'edit_file.html', {
        'file': file_obj,
        'form': form,
        'extension': extension,
        'editable_text': editable_text,
        'text_preview': text_preview,
        'image_preview': image_preview,
        'pdf_preview': pdf_preview,
        'ChangeTypes': ChangeTypes,
    })


# -------------------------
# Delete
# -------------------------
@login_required
def file_delete(request, pk):
    file_obj = get_object_or_404(UploadedFile, pk=pk)
    if request.user != file_obj.owner:
        return HttpResponseForbidden("You are not allowed to delete this file.")

    # delete converted PDF if exists
    try:
        if file_obj.converted and file_obj.converted.path and os.path.exists(file_obj.converted.path):
            os.remove(file_obj.converted.path)
    except Exception:
        pass

    # delete original file
    try:
        if os.path.exists(file_obj.file.path):
            os.remove(file_obj.file.path)
    except Exception:
        pass

    file_obj.delete()
    messages.success(request, "File deleted.")
    return redirect('file_list')


# -------------------------
# Convert to PDF
# -------------------------
def _convert_with_libreoffice(file_obj):
    """
    Shared helper that runs LibreOffice headless conversion and updates the model.
    Returns (success: bool, message: str)
    """
    input_path = file_obj.file.path
    output_dir = os.path.join(settings.MEDIA_ROOT, 'converted')
    os.makedirs(output_dir, exist_ok=True)

    print(f"[DEBUG] Converting file: {input_path}")
    print(f"[DEBUG] Output directory: {output_dir}")

    try:
        result = subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            input_path, '--outdir', output_dir
        ], check=True, capture_output=True, timeout=120)

        print(f"[DEBUG] LibreOffice stdout: {result.stdout.decode(errors='ignore')}")
        print(f"[DEBUG] LibreOffice stderr: {result.stderr.decode(errors='ignore')}")

    except subprocess.TimeoutExpired:
        return False, "Conversion timed out after 120 seconds."
    except subprocess.CalledProcessError as e:
        error_output = e.stderr.decode(errors='ignore') if e.stderr else 'No error output'
        print(f"[DEBUG] LibreOffice error: {error_output}")
        return False, f"Conversion failed with error code {e.returncode}"
    except Exception as e:
        print(f"[DEBUG] Exception: {str(e)}")
        return False, f"Conversion failed: {e}"

    base_name = os.path.splitext(os.path.basename(input_path))[0]
    converted_name = f"{base_name}.pdf"
    converted_full = os.path.join(output_dir, converted_name)

    print(f"[DEBUG] Looking for converted file: {converted_full}")
    print(f"[DEBUG] File exists: {os.path.exists(converted_full)}")

    if not os.path.exists(converted_full):
        return False, f"Conversion did not produce a PDF at: {converted_full}"

    file_size = os.path.getsize(converted_full)
    print(f"[DEBUG] Converted PDF size: {file_size} bytes")

    if file_size == 0:
        return False, "Conversion produced an empty PDF file."

    try:
        # Remove old converted file if it exists
        if file_obj.converted:
            try:
                old_path = file_obj.converted.path
                if old_path and os.path.exists(old_path):
                    os.remove(old_path)
            except Exception:
                pass

        with open(converted_full, 'rb') as f:
            file_obj.converted.save(converted_name, File(f), save=False)

        file_obj.file_name_if_converted = converted_name
        file_obj.save()

        print(f"[DEBUG] Model updated successfully")
        print(f"[DEBUG] file_obj.converted.name: {file_obj.converted.name}")
        print(f"[DEBUG] file_obj.converted.path: {file_obj.converted.path}")

    except Exception as e:
        print(f"[DEBUG] Model save exception: {str(e)}")
        return False, f"Failed to save converted file to database: {e}"

    return True, f"Converted to PDF successfully! ({file_size} bytes)"


@login_required
def convert_to_pdf(request, pk):
    file_obj = get_object_or_404(UploadedFile, pk=pk)
    if request.user != file_obj.owner:
        return HttpResponseForbidden("No permission to convert.")

    success, feedback = _convert_with_libreoffice(file_obj)
    if success:
        messages.success(request, feedback)
    else:
        messages.error(request, feedback)
    return redirect('file_detail', pk=pk)


@login_required
def update_file_status(request, pk, action):
    file_obj = get_object_or_404(UploadedFile, pk=pk)
    if not is_program_super_user(request.user):
        return HttpResponseForbidden("Only program super users can update status.")
    if request.method != 'POST':
        raise Http404("Invalid method")

    status_map = {
        'start': FileStatus.IN_REVIEW,
        'approve': FileStatus.APPROVED,
        'reject': FileStatus.REJECTED,
    }
    if action not in status_map:
        raise Http404("Unknown action")

    new_status = status_map[action]
    file_obj.status = new_status
    file_obj.reviewed_by = request.user
    file_obj.reviewed_at = timezone.now()
    file_obj.save()

    if action == 'approve':
        notify_users(
            User.objects.filter(is_active=True),
            request.user,
            Notification.Types.FILE_APPROVED,
            f"{file_obj.filename} has been approved.",
            file_obj,
        )
        messages.success(request, "File approved.")
    elif action == 'reject':
        notify_users(
            User.objects.filter(id=file_obj.owner_id),
            request.user,
            Notification.Types.FILE_REJECTED,
            f"{file_obj.filename} was rejected.",
            file_obj,
        )
        messages.warning(request, "File rejected.")
    else:
        messages.info(request, "File moved to in-review.")

    return redirect('file_detail', pk=pk)


@login_required
def notifications_list(request):
    notifications = request.user.notifications.select_related('sender', 'related_file').order_by('-created_at')

    if request.method == 'POST':
        action = request.POST.get('action')
        notification_id = request.POST.get('notification_id')

        if action == 'read_all':
            count = notifications.filter(is_read=False).update(is_read=True)
            if count:
                messages.success(request, f"Marked {count} notification(s) as read.")
            else:
                messages.info(request, "No unread notifications.")
            return redirect('notifications')

        if not notification_id:
            messages.error(request, "Invalid notification request.")
            return redirect('notifications')

        notif = get_object_or_404(Notification, pk=notification_id, recipient=request.user)

        if action == 'mark_read':
            if not notif.is_read:
                notif.is_read = True
                notif.save(update_fields=['is_read'])
            messages.success(request, "Notification marked as read.")
        elif action == 'mark_unread':
            if notif.is_read:
                notif.is_read = False
                notif.save(update_fields=['is_read'])
            messages.success(request, "Notification marked as unread.")
        elif action == 'dismiss':
            notif.delete()
            messages.success(request, "Notification dismissed.")
        else:
            messages.error(request, "Unknown notification action.")

        return redirect('notifications')

    return render(request, 'notifications.html', {'notifications': notifications})


# -------------------------
# Add comment (separate view)
# -------------------------
@login_required
def add_comment(request, pk):
    file_obj = get_object_or_404(UploadedFile, pk=pk)
    if request.user != file_obj.owner:
        return HttpResponseForbidden("Only owner can comment on their files.")

    if request.method != 'POST':
        raise Http404("Invalid method")

    form = CommentForm(request.POST)
    if form.is_valid():
        comment = form.save(commit=False)
        comment.file = file_obj
        comment.user = request.user
        comment.save()
        messages.success(request, "Comment added.")
    else:
        messages.error(request, "Comment failed.")

    return redirect('file_detail', pk=pk)